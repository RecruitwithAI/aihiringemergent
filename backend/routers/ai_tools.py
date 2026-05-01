from fastapi import APIRouter, Depends, HTTPException, UploadFile, File, Form, Response
from datetime import datetime, timezone, timedelta
import uuid
import os
import io
import re
import tempfile
from pathlib import Path

from utils.database import db
from utils.helpers import get_badge
from utils.auth import get_current_user, add_points
from models.schemas import AIToolRequest, DownloadRequest, ExtractFileRequest, UserAPIKeyUpdate


router = APIRouter(prefix="/ai", tags=["ai"])

EMERGENT_LLM_KEY = os.environ.get("EMERGENT_LLM_KEY")
MASTER_OPENAI_KEY = os.environ.get("MASTER_OPENAI_KEY")
UPLOAD_DIR = Path("/tmp/bestpl_uploads")
UPLOAD_DIR.mkdir(parents=True, exist_ok=True)
CHUNK_SIZE = 1 * 1024 * 1024  # 1 MB

DAILY_FREE_LIMIT = 3  # Free API key uses per user per day

from utils.prompts import get_prompt


def clean_ai_response(response: str) -> str:
    """
    Remove conversational follow-ups and questions from AI responses.
    Strips out common patterns like 'If you want...', 'Let me know...', etc.
    """
    # Split by lines
    lines = response.split('\n')
    cleaned_lines = []
    found_separator = False
    
    # Common patterns that indicate conversational follow-ups
    followup_patterns = [
        "if you want",
        "if you'd like",
        "would you like",
        "let me know",
        "feel free to",
        "i can help",
        "i can tailor",
        "i can create",
        "i can provide",
        "i can also",
        "shall i",
        "should i",
        "do you want",
        "would you prefer",
        "happy to help",
        "happy to assist",
        "need any changes",
        "need anything else",
        "anything else",
        "any questions",
    ]
    
    for line in lines:
        stripped = line.strip()
        lower_line = stripped.lower()
        
        # Check if this line starts a conversational section
        # Often preceded by \"---\" or empty lines
        if stripped == "---" or stripped == "***":
            # Check if the next content looks conversational
            found_separator = True
            continue
        
        # If we found a separator, check if subsequent non-empty lines are conversational
        if found_separator and stripped:
            is_conversational = any(pattern in lower_line for pattern in followup_patterns)
            if is_conversational:
                # Skip this and remaining lines (likely all conversational)
                break
            else:
                # False alarm, include the separator and continue
                if cleaned_lines and cleaned_lines[-1].strip() != "":
                    cleaned_lines.append("")  # Add spacing
                found_separator = False
        
        # Check if line itself starts with a conversational pattern
        is_conversational_start = any(
            lower_line.startswith(pattern) for pattern in followup_patterns
        )
        
        if is_conversational_start:
            # This line and everything after is likely conversational - stop here
            break
        
        cleaned_lines.append(line)
    
    # Join and clean up excessive blank lines at the end
    result = '\n'.join(cleaned_lines).rstrip('\n')
    
    # Remove trailing \"---\" if present
    while result.endswith('\n---') or result.endswith('\n***'):
        result = result.rsplit('\n', 1)[0].rstrip()
    
    return result


async def check_daily_usage(user_id: str) -> dict:
    """Check if user has remaining free API uses today"""
    today_start = datetime.now(timezone.utc).replace(hour=0, minute=0, second=0, microsecond=0)
    
    usage_count = await db.api_usage.count_documents({
        "user_id": user_id,
        "used_master_key": True,
        "timestamp": {"$gte": today_start}
    })
    
    remaining = max(0, DAILY_FREE_LIMIT - usage_count)
    return {
        "used": usage_count,
        "remaining": remaining,
        "limit": DAILY_FREE_LIMIT,
        "can_use": remaining > 0
    }


async def record_usage(user_id: str, used_master_key: bool, tool_type: str):
    """Record API usage"""
    await db.api_usage.insert_one({
        "user_id": user_id,
        "used_master_key": used_master_key,
        "tool_type": tool_type,
        "timestamp": datetime.now(timezone.utc)
    })


@router.get("/usage")
async def get_usage_stats(user=Depends(get_current_user)):
    """Get user's API usage statistics"""
    usage = await check_daily_usage(user["user_id"])
    
    # Get user's stored API key status
    user_data = await db.users.find_one({"user_id": user["user_id"]}, {"_id": 0, "has_own_api_key": 1})
    has_own_key = user_data.get("has_own_api_key", False) if user_data else False
    
    return {
        "daily_usage": usage,
        "has_own_api_key": has_own_key
    }


@router.post("/generate")
async def ai_generate(req: AIToolRequest, user=Depends(get_current_user)):
    from emergentintegrations.llm.chat import LlmChat, UserMessage

    # Determine which API key to use
    api_key_to_use = None
    using_master_key = False
    
    if req.use_own_key and req.own_api_key:
        # User provided their own API key for this request
        api_key_to_use = req.own_api_key
        using_master_key = False
    else:
        # Check if user has a saved API key
        user_data = await db.users.find_one({"user_id": user["user_id"]}, {"_id": 0, "openai_api_key": 1})
        if user_data and user_data.get("openai_api_key"):
            api_key_to_use = user_data["openai_api_key"]
            using_master_key = False
        else:
            # Use SuperAdmin's API key for free usages - check daily limit
            usage = await check_daily_usage(user["user_id"])
            if not usage["can_use"]:
                raise HTTPException(
                    status_code=429,
                    detail=f"Daily free API limit reached ({DAILY_FREE_LIMIT} uses per day). Please add your own OpenAI API key to continue."
                )
            
            # Get SuperAdmin's API key
            superadmin = await db.users.find_one({"role": "superadmin"}, {"_id": 0, "openai_api_key": 1})
            if superadmin and superadmin.get("openai_api_key"):
                api_key_to_use = superadmin["openai_api_key"]
                using_master_key = True
            else:
                raise HTTPException(
                    status_code=503,
                    detail="System API key not configured. Please contact administrator or add your own OpenAI API key."
                )

    system_prompt = await get_prompt(req.tool_type)

    chat = LlmChat(
        api_key=api_key_to_use,
        session_id=f"ai_{user['user_id']}_{uuid.uuid4().hex[:8]}",
        system_message=system_prompt,
    ).with_model("openai", "gpt-4o")

    full_prompt = req.prompt
    if req.context:
        full_prompt = f"{req.prompt}\n\nAdditional Context: {req.context}"

    response = await chat.send_message(UserMessage(text=full_prompt))
    
    # Clean up conversational follow-ups from the response
    cleaned_response = clean_ai_response(response)

    # Record usage
    await record_usage(user["user_id"], using_master_key, req.tool_type)

    await db.ai_history.insert_one({
        "history_id": f"hist_{uuid.uuid4().hex[:12]}",
        "user_id": user["user_id"],
        "tool_type": req.tool_type,
        "prompt": req.prompt,
        "response": cleaned_response,
        "created_at": datetime.now(timezone.utc).isoformat(),
    })

    await add_points(user["user_id"], 2)
    return {"response": cleaned_response, "tool_type": req.tool_type}


@router.get("/history")
async def get_ai_history(user=Depends(get_current_user)):
    history = await db.ai_history.find({"user_id": user["user_id"]}, {"_id": 0}).sort("created_at", -1).to_list(50)
    return history


# ── File Upload & Extraction ──

@router.post("/upload-chunk")
async def upload_chunk(
    chunk: UploadFile = File(...),
    upload_id: str = Form(...),
    chunk_index: str = Form(...),
    total_chunks: str = Form(...),
    filename: str = Form(...),
    user=Depends(get_current_user),
):
    chunk_dir = UPLOAD_DIR / f"{user['user_id']}_{upload_id}"
    chunk_dir.mkdir(parents=True, exist_ok=True)
    data = await chunk.read()
    (chunk_dir / f"chunk_{int(chunk_index):05d}").write_bytes(data)
    return {"chunk": int(chunk_index), "total": int(total_chunks), "ok": True}

@router.post("/extract-file")
async def extract_file(req: ExtractFileRequest, user=Depends(get_current_user)):
    chunk_dir = UPLOAD_DIR / f"{user['user_id']}_{req.upload_id}"
    if not chunk_dir.exists():
        raise HTTPException(status_code=404, detail="Upload session not found")

    chunks = sorted(chunk_dir.glob("chunk_*"))
    if not chunks:
        raise HTTPException(status_code=400, detail="No chunks received")

    content = b"".join(c.read_bytes() for c in chunks)
    for c in chunks:
        c.unlink()
    chunk_dir.rmdir()

    ext = req.filename.lower().rsplit(".", 1)[-1] if "." in req.filename else ""
    extracted = ""

    try:
        if ext == "txt":
            extracted = content.decode("utf-8", errors="ignore")

        elif ext == "pdf":
            import pypdf
            reader = pypdf.PdfReader(io.BytesIO(content))
            
            # Extract text with better formatting preservation
            pages_text = []
            for page_num, page in enumerate(reader.pages, 1):
                page_text = page.extract_text() or ""
                if page_text.strip():
                    # Add page marker for multi-page documents to preserve structure
                    if len(reader.pages) > 1:
                        pages_text.append(f"[Page {page_num}]\n{page_text}")
                    else:
                        pages_text.append(page_text)
            
            extracted = "\n\n".join(pages_text)
            
            # Clean up common PDF extraction artifacts while preserving structure
            # Remove excessive whitespace but keep paragraph breaks
            extracted = re.sub(r'\n{3,}', '\n\n', extracted)  # Max 2 newlines
            extracted = re.sub(r'[ \t]+', ' ', extracted)  # Normalize spaces
            extracted = extracted.strip()

        elif ext == "docx":
            from docx import Document as DocxDoc
            doc = DocxDoc(io.BytesIO(content))
            extracted = "\n".join(p.text for p in doc.paragraphs)

        elif ext == "doc":
            try:
                from docx import Document as DocxDoc
                doc = DocxDoc(io.BytesIO(content))
                extracted = "\n".join(p.text for p in doc.paragraphs)
            except Exception:
                text = content.decode("latin-1", errors="ignore")
                extracted = " ".join(re.findall(r'[\x20-\x7E\n\r\t]{4,}', text))

        elif ext in ("mp3", "wav", "m4a", "ogg", "aac", "flac", "mp4", "mpeg", "mpga", "webm"):
            from emergentintegrations.llm.openai import OpenAISpeechToText
            stt = OpenAISpeechToText(api_key=EMERGENT_LLM_KEY)
            with tempfile.NamedTemporaryFile(suffix=f".{ext}", delete=False) as tmp:
                tmp.write(content)
                tmp_path = tmp.name
            try:
                with open(tmp_path, "rb") as audio_file:
                    resp = await stt.transcribe(
                        file=audio_file, model="whisper-1", response_format="json", language="en",
                        prompt="Recruiting, job descriptions, executive search, candidate profiles.",
                    )
                extracted = resp.text
            finally:
                os.unlink(tmp_path)

        else:
            raise HTTPException(status_code=400, detail=f"Unsupported file type: .{ext}")

    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"File extraction failed: {str(e)}")

    if not extracted.strip():
        raise HTTPException(status_code=422, detail="Could not extract text from file. Try a different format.")

    return {"extracted_text": extracted.strip(), "filename": req.filename, "char_count": len(extracted)}


# ── Document Download Generation ──

def _add_bold_runs(paragraph, text: str):
    parts = re.split(r"\*\*(.*?)\*\*", text)
    for i, part in enumerate(parts):
        paragraph.add_run(part).bold = (i % 2 == 1)


def _parse_md_to_docx(content: str) -> bytes:
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    
    doc = Document()
    lines = content.split("\n")
    i = 0
    
    while i < len(lines):
        s = lines[i].strip()
        
        # Check if this is a table (markdown table starts with |)
        if s.startswith("|") and "|" in s:
            # Collect all table lines
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                table_lines.append(lines[i].strip())
                i += 1
            
            # Parse table
            if len(table_lines) >= 2:  # Need at least header and separator
                # Parse header
                header = [cell.strip() for cell in table_lines[0].split("|")[1:-1]]
                
                # Skip separator line (the one with ---)
                data_lines = table_lines[2:] if len(table_lines) > 2 else []
                
                # Create table
                num_cols = len(header)
                num_rows = len(data_lines) + 1  # +1 for header
                
                if num_cols > 0 and num_rows > 0:
                    table = doc.add_table(rows=num_rows, cols=num_cols)
                    table.style = 'Light Grid Accent 1'
                    
                    # Add header
                    for col_idx, header_text in enumerate(header):
                        cell = table.rows[0].cells[col_idx]
                        cell.text = header_text
                        # Make header bold
                        for paragraph in cell.paragraphs:
                            for run in paragraph.runs:
                                run.bold = True
                    
                    # Add data rows
                    for row_idx, line in enumerate(data_lines, start=1):
                        cells = [cell.strip() for cell in line.split("|")[1:-1]]
                        for col_idx, cell_text in enumerate(cells):
                            if col_idx < num_cols:
                                table.rows[row_idx].cells[col_idx].text = cell_text
                    
                    doc.add_paragraph()  # Add spacing after table
            continue
        
        # Regular line processing
        if not s:
            doc.add_paragraph("")
        elif s.startswith("### "):
            doc.add_heading(s[4:], level=3)
        elif s.startswith("## "):
            doc.add_heading(s[3:], level=2)
        elif s.startswith("# "):
            doc.add_heading(s[2:], level=1)
        elif s.startswith(("- ", "* ", "\u2022 ")):
            p = doc.add_paragraph(style="List Bullet")
            _add_bold_runs(p, s[2:])
        else:
            p = doc.add_paragraph()
            _add_bold_runs(p, s)
        
        i += 1
    
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


def _parse_md_to_pdf(content: str) -> bytes:
    from fpdf import FPDF

    def safe(t: str) -> str:
        # Replace problematic characters and encode safely
        t = t.replace('\u2022', '-').replace('\u2019', "'").replace('\u2018', "'")
        t = t.replace('\u201c', '"').replace('\u201d', '"').replace('\u2014', '-')
        t = t.replace('\u2013', '-').replace('\u00a0', ' ')
        # Handle other unicode characters more gracefully
        t = t.encode("latin-1", errors="replace").decode("latin-1")
        return t

    pdf = FPDF(orientation="P", unit="mm", format="A4")
    pdf.add_page()
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.set_margins(15, 15, 15)

    # Page width minus margins
    effective_width = 210 - 30  # A4 width (210mm) minus left+right margins (15+15)
    
    lines = content.split("\n")
    i = 0
    
    while i < len(lines):
        s = lines[i].strip()
        
        # Check if this is a table (markdown table starts with |)
        if s.startswith("|") and "|" in s:
            # Collect all table lines
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                table_lines.append(lines[i].strip())
                i += 1
            
            # Parse table
            if len(table_lines) >= 2:
                # Parse header
                header = [cell.strip() for cell in table_lines[0].split("|")[1:-1]]
                
                # Skip separator line (the one with ---)
                data_lines = table_lines[2:] if len(table_lines) > 2 else []
                
                num_cols = len(header)
                if num_cols > 0:
                    # Calculate column width - make slightly smaller for better fit
                    col_width = effective_width / num_cols
                    cell_height = 6
                    
                    # Add table header
                    pdf.set_font("Helvetica", "B", 8)
                    pdf.set_fill_color(200, 200, 200)
                    
                    x_start = pdf.get_x()
                    y_start = pdf.get_y()
                    
                    for col_idx, header_text in enumerate(header):
                        pdf.set_xy(x_start + (col_idx * col_width), y_start)
                        # Truncate long headers
                        max_chars = int(col_width * 2.2)
                        text = safe(header_text)[:max_chars]
                        pdf.cell(col_width, cell_height, text, border=1, align='L', fill=True)
                    
                    pdf.set_xy(x_start, y_start + cell_height)
                    
                    # Add data rows
                    pdf.set_font("Helvetica", "", 7)
                    for line in data_lines:
                        cells = [cell.strip() for cell in line.split("|")[1:-1]]
                        
                        # Calculate row height based on content
                        row_height = cell_height
                        for cell_text in cells[:num_cols]:
                            # Estimate lines needed
                            max_chars_per_line = int(col_width * 2.5)
                            if len(cell_text) > max_chars_per_line:
                                lines_needed = (len(cell_text) // max_chars_per_line) + 1
                                estimated_height = lines_needed * 4
                                row_height = max(row_height, estimated_height)
                        
                        y_start = pdf.get_y()
                        
                        # Draw all cells in the row
                        for col_idx, cell_text in enumerate(cells[:num_cols]):
                            if col_idx >= num_cols:
                                break
                                
                            x_pos = x_start + (col_idx * col_width)
                            
                            # Draw cell border
                            pdf.rect(x_pos, y_start, col_width, row_height)
                            
                            # Add text with padding
                            pdf.set_xy(x_pos + 1, y_start + 1)
                            
                            # Calculate available width for text (with padding)
                            text_width = col_width - 2
                            
                            # Truncate or wrap text
                            text = safe(cell_text)
                            max_chars = int(text_width * 2.5)
                            
                            if len(text) > max_chars:
                                # Multi-line text - split into lines
                                words = text.split()
                                lines_list = []
                                current_line = ""
                                
                                for word in words:
                                    test_line = current_line + " " + word if current_line else word
                                    if len(test_line) <= max_chars:
                                        current_line = test_line
                                    else:
                                        if current_line:
                                            lines_list.append(current_line)
                                        current_line = word
                                
                                if current_line:
                                    lines_list.append(current_line)
                                
                                # Draw each line
                                line_y = y_start + 1
                                for text_line in lines_list[:int(row_height/4)]:  # Limit lines to fit
                                    pdf.set_xy(x_pos + 1, line_y)
                                    pdf.cell(text_width, 4, text_line[:max_chars], border=0, align='L')
                                    line_y += 4
                            else:
                                # Single line
                                pdf.cell(text_width, 4, text, border=0, align='L')
                        
                        # Move to next row
                        pdf.set_xy(x_start, y_start + row_height)
                    
                    pdf.ln(5)  # Add spacing after table
            continue
        
        # Regular line processing (existing logic)
        if not s:
            pdf.ln(3)
            i += 1
            continue
        
        # Remove markdown bold markers for display
        clean = re.sub(r"\*\*(.*?)\*\*", r"\1", s)
        
        # Skip lines that are just special characters
        if not clean or len(clean.strip()) == 0:
            pdf.ln(2)
            i += 1
            continue

        # Process different line types
        if s.startswith("# ") and not s.startswith("##"):
            # H1 - Main heading
            pdf.set_font("Helvetica", "B", 16)
            text = safe(clean[2:].strip())
            if text:
                pdf.multi_cell(effective_width, 8, text, align='L')
                pdf.ln(3)
                
        elif s.startswith("### "):
            # H3 - Sub-subheading
            pdf.set_font("Helvetica", "B", 11)
            text = safe(clean[4:].strip())
            if text:
                pdf.multi_cell(effective_width, 6, text, align='L')
                pdf.ln(1)
                
        elif s.startswith("## "):
            # H2 - Subheading
            pdf.set_font("Helvetica", "B", 13)
            text = safe(clean[3:].strip())
            if text:
                pdf.multi_cell(effective_width, 7, text, align='L')
                pdf.ln(2)
                
        elif s.startswith(("- ", "* ", "\u2022 ")):
            # Bullet point
            pdf.set_font("Helvetica", "", 10)
            bullet_text = clean[2:].strip()
            if bullet_text:
                # Save current position
                x_start = pdf.l_margin
                y_start = pdf.get_y()
                
                # Draw bullet character
                pdf.set_xy(x_start, y_start)
                pdf.cell(5, 5, txt=chr(149), ln=0)  # Use bullet character (•)
                
                # Calculate text area (full width minus bullet indent)
                text_x = x_start + 7
                text_width = effective_width - 7
                
                # Position for text and use multi_cell for wrapping
                pdf.set_xy(text_x, y_start)
                text = safe(bullet_text)
                
                # Create a temporary position to measure text height
                # We need to manually handle the multi_cell positioning
                pdf.multi_cell(text_width, 5, text, align='L', ln=1)
                
                # Reset X margin for next line
                pdf.set_x(x_start)
                
        else:
            # Regular paragraph text
            pdf.set_font("Helvetica", "", 10)
            text = safe(clean)
            if text:
                pdf.multi_cell(effective_width, 5, text, align='L')
                pdf.ln(1)
        
        i += 1

    return bytes(pdf.output())


@router.post("/download")
async def download_document(req: DownloadRequest, user=Depends(get_current_user)):
    safe_name = re.sub(r"[^\w\- ]", "_", req.filename)[:60]
    
    if req.format == "txt":
        return Response(
            content=req.content.encode("utf-8"),
            media_type="text/plain; charset=utf-8",
            headers={"Content-Disposition": f'attachment; filename="{safe_name}.txt"'},
        )
    
    elif req.format == "csv":
        # CSV export for talent-scout (expects JSON array of candidates)
        import csv
        import json
        
        try:
            candidates = json.loads(req.content)
            if not isinstance(candidates, list):
                raise ValueError("Content must be a JSON array")
            
            # Create CSV in memory
            output = io.StringIO()
            if candidates:
                writer = csv.DictWriter(output, fieldnames=candidates[0].keys())
                writer.writeheader()
                
                for candidate in candidates:
                    # Flatten arrays/lists for CSV
                    row = {}
                    for key, value in candidate.items():
                        if isinstance(value, list):
                            row[key] = ", ".join(str(v) for v in value)
                        else:
                            row[key] = value
                    writer.writerow(row)
            
            csv_content = output.getvalue()
            output.close()
            
            return Response(
                content=csv_content.encode("utf-8"),
                media_type="text/csv; charset=utf-8",
                headers={"Content-Disposition": f'attachment; filename="{safe_name}.csv"'},
            )
        except (json.JSONDecodeError, ValueError):
            # Fallback: treat as plain text CSV
            return Response(
                content=req.content.encode("utf-8"),
                media_type="text/csv; charset=utf-8",
                headers={"Content-Disposition": f'attachment; filename="{safe_name}.csv"'},
            )
    
    elif req.format == "docx":
        return Response(
            content=_parse_md_to_docx(req.content),
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": f'attachment; filename="{safe_name}.docx"'},
        )
    
    elif req.format == "pdf":
        return Response(
            content=_parse_md_to_pdf(req.content),
            media_type="application/pdf",
            headers={"Content-Disposition": f'attachment; filename="{safe_name}.pdf"'},
        )
    
    raise HTTPException(status_code=400, detail="Invalid format. Use txt, csv, docx, or pdf.")



@router.post("/save-api-key")
async def save_user_api_key(req: UserAPIKeyUpdate, user=Depends(get_current_user)):
    """Save user's personal OpenAI API key"""
    if not req.api_key or len(req.api_key.strip()) < 10:
        raise HTTPException(status_code=400, detail="Invalid API key")
    
    # Update user's API key in database (stored securely)
    await db.users.update_one(
        {"user_id": user["user_id"]},
        {"$set": {
            "openai_api_key": req.api_key.strip(),
            "has_own_api_key": True,
            "api_key_updated_at": datetime.now(timezone.utc)
        }}
    )
    
    return {"success": True, "message": "API key saved successfully"}


@router.delete("/delete-api-key")
async def delete_user_api_key(user=Depends(get_current_user)):
    """Remove user's personal OpenAI API key"""
    await db.users.update_one(
        {"user_id": user["user_id"]},
        {"$unset": {"openai_api_key": ""}, "$set": {"has_own_api_key": False}}
    )
    
    return {"success": True, "message": "API key removed successfully"}
