"""
Markdown → DOCX / PDF export, decomposed into small focused helpers.

Public API (used by routers/ai_tools.py :: download endpoint):
  markdown_to_docx(content) -> bytes
  markdown_to_pdf(content)  -> bytes

Behavior matches the original inline implementation exactly (verified e2e):
headings (#/##/###), bullets, **bold** runs, and markdown tables.
"""
import io
import re


# ============================================================
# Shared markdown helpers
# ============================================================

def _collect_table_block(lines, i):
    """Collect consecutive markdown table lines starting at index i.

    Returns (table_lines, next_index)."""
    table_lines = []
    while i < len(lines) and lines[i].strip().startswith("|"):
        table_lines.append(lines[i].strip())
        i += 1
    return table_lines, i


def _split_table_row(line):
    """'| a | b |' -> ['a', 'b']"""
    return [cell.strip() for cell in line.split("|")[1:-1]]


def _parse_table(table_lines):
    """Returns (header_cells, data_lines) or (None, None) if not a valid table."""
    if len(table_lines) < 2:  # Need at least header and separator
        return None, None
    header = _split_table_row(table_lines[0])
    # Skip separator line (the one with ---)
    data_lines = table_lines[2:] if len(table_lines) > 2 else []
    return header, data_lines


# ============================================================
# DOCX export
# ============================================================

def _add_bold_runs(paragraph, text: str):
    parts = re.split(r"\*\*(.*?)\*\*", text)
    for i, part in enumerate(parts):
        paragraph.add_run(part).bold = (i % 2 == 1)


def _add_docx_table(doc, table_lines):
    header, data_lines = _parse_table(table_lines)
    if header is None:
        return

    num_cols = len(header)
    num_rows = len(data_lines) + 1  # +1 for header
    if num_cols == 0 or num_rows == 0:
        return

    table = doc.add_table(rows=num_rows, cols=num_cols)
    table.style = 'Light Grid Accent 1'

    # Header row (bold)
    for col_idx, header_text in enumerate(header):
        cell = table.rows[0].cells[col_idx]
        cell.text = header_text
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True

    # Data rows
    for row_idx, line in enumerate(data_lines, start=1):
        cells = _split_table_row(line)
        for col_idx, cell_text in enumerate(cells):
            if col_idx < num_cols:
                table.rows[row_idx].cells[col_idx].text = cell_text

    doc.add_paragraph()  # Spacing after table


def _add_docx_line(doc, s: str):
    """Render a single non-table markdown line into the document."""
    if not s:
        doc.add_paragraph("")
    elif s.startswith("### "):
        doc.add_heading(s[4:], level=3)
    elif s.startswith("## "):
        doc.add_heading(s[3:], level=2)
    elif s.startswith("# "):
        doc.add_heading(s[2:], level=1)
    elif s.startswith(("- ", "* ", "\u2022 ")):
        p = doc.add_paragraph(style="List Bullet")
        _add_bold_runs(p, s[2:])
    else:
        p = doc.add_paragraph()
        _add_bold_runs(p, s)


def markdown_to_docx(content: str) -> bytes:
    from docx import Document

    doc = Document()
    lines = content.split("\n")
    i = 0
    while i < len(lines):
        s = lines[i].strip()
        if s.startswith("|"):
            table_lines, i = _collect_table_block(lines, i)
            _add_docx_table(doc, table_lines)
            continue
        _add_docx_line(doc, s)
        i += 1

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


# ============================================================
# PDF export
# ============================================================

_PDF_EFFECTIVE_WIDTH = 210 - 30  # A4 width (210mm) minus left+right margins (15+15)


def _pdf_safe(t: str) -> str:
    """Replace problematic characters and encode latin-1 safely."""
    t = t.replace('\u2022', '-').replace('\u2019', "'").replace('\u2018', "'")
    t = t.replace('\u201c', '"').replace('\u201d', '"').replace('\u2014', '-')
    t = t.replace('\u2013', '-').replace('\u00a0', ' ')
    return t.encode("latin-1", errors="replace").decode("latin-1")


def _wrap_words(text: str, max_chars: int):
    """Greedy word-wrap into lines of at most max_chars."""
    words = text.split()
    lines_list = []
    current_line = ""
    for word in words:
        test_line = current_line + " " + word if current_line else word
        if len(test_line) <= max_chars:
            current_line = test_line
        else:
            if current_line:
                lines_list.append(current_line)
            current_line = word
    if current_line:
        lines_list.append(current_line)
    return lines_list


def _estimate_pdf_row_height(cells, num_cols, col_width, base_height):
    """Estimate row height from the longest cell content."""
    row_height = base_height
    max_chars_per_line = int(col_width * 2.5)
    for cell_text in cells[:num_cols]:
        if len(cell_text) > max_chars_per_line:
            lines_needed = (len(cell_text) // max_chars_per_line) + 1
            row_height = max(row_height, lines_needed * 4)
    return row_height


def _render_pdf_cell(pdf, cell_text, x_pos, y_start, col_width, row_height):
    """Render one bordered table cell with wrapped/truncated text."""
    pdf.rect(x_pos, y_start, col_width, row_height)
    pdf.set_xy(x_pos + 1, y_start + 1)

    text_width = col_width - 2
    text = _pdf_safe(cell_text)
    max_chars = int(text_width * 2.5)

    if len(text) <= max_chars:
        pdf.cell(text_width, 4, text, border=0, align='L')
        return

    # Multi-line: wrap and draw each line, limited to fit row height
    line_y = y_start + 1
    for text_line in _wrap_words(text, max_chars)[:int(row_height / 4)]:
        pdf.set_xy(x_pos + 1, line_y)
        pdf.cell(text_width, 4, text_line[:max_chars], border=0, align='L')
        line_y += 4


def _render_pdf_table(pdf, table_lines):
    header, data_lines = _parse_table(table_lines)
    if header is None:
        return
    num_cols = len(header)
    if num_cols == 0:
        return

    col_width = _PDF_EFFECTIVE_WIDTH / num_cols
    cell_height = 6

    # Header row (bold, filled)
    pdf.set_font("Helvetica", "B", 8)
    pdf.set_fill_color(200, 200, 200)
    x_start = pdf.get_x()
    y_start = pdf.get_y()
    for col_idx, header_text in enumerate(header):
        pdf.set_xy(x_start + (col_idx * col_width), y_start)
        max_chars = int(col_width * 2.2)
        text = _pdf_safe(header_text)[:max_chars]
        pdf.cell(col_width, cell_height, text, border=1, align='L', fill=True)
    pdf.set_xy(x_start, y_start + cell_height)

    # Data rows
    pdf.set_font("Helvetica", "", 7)
    for line in data_lines:
        cells = _split_table_row(line)
        row_height = _estimate_pdf_row_height(cells, num_cols, col_width, cell_height)
        y_start = pdf.get_y()
        for col_idx, cell_text in enumerate(cells[:num_cols]):
            _render_pdf_cell(pdf, cell_text, x_start + (col_idx * col_width),
                             y_start, col_width, row_height)
        pdf.set_xy(x_start, y_start + row_height)

    pdf.ln(5)  # Spacing after table


def _render_pdf_heading(pdf, text, font_size, line_height, spacing_after):
    pdf.set_font("Helvetica", "B", font_size)
    text = _pdf_safe(text.strip())
    if text:
        pdf.multi_cell(_PDF_EFFECTIVE_WIDTH, line_height, text, align='L')
        pdf.ln(spacing_after)


def _render_pdf_bullet(pdf, bullet_text):
    pdf.set_font("Helvetica", "", 10)
    bullet_text = bullet_text.strip()
    if not bullet_text:
        return
    x_start = pdf.l_margin
    y_start = pdf.get_y()

    # Draw bullet character
    pdf.set_xy(x_start, y_start)
    pdf.cell(5, 5, txt=chr(149), ln=0)

    # Wrapped text to the right of the bullet
    text_x = x_start + 7
    text_width = _PDF_EFFECTIVE_WIDTH - 7
    pdf.set_xy(text_x, y_start)
    pdf.multi_cell(text_width, 5, _pdf_safe(bullet_text), align='L', ln=1)
    pdf.set_x(x_start)


def _render_pdf_line(pdf, s: str):
    """Render a single non-table markdown line. Early-returns per line type."""
    if not s:
        pdf.ln(3)
        return

    # Remove markdown bold markers for display
    clean = re.sub(r"\*\*(.*?)\*\*", r"\1", s)
    if not clean.strip():
        pdf.ln(2)
        return

    if s.startswith("# ") and not s.startswith("##"):
        _render_pdf_heading(pdf, clean[2:], font_size=16, line_height=8, spacing_after=3)
        return
    if s.startswith("### "):
        _render_pdf_heading(pdf, clean[4:], font_size=11, line_height=6, spacing_after=1)
        return
    if s.startswith("## "):
        _render_pdf_heading(pdf, clean[3:], font_size=13, line_height=7, spacing_after=2)
        return
    if s.startswith(("- ", "* ", "\u2022 ")):
        _render_pdf_bullet(pdf, clean[2:])
        return

    # Regular paragraph text
    pdf.set_font("Helvetica", "", 10)
    text = _pdf_safe(clean)
    if text:
        pdf.multi_cell(_PDF_EFFECTIVE_WIDTH, 5, text, align='L')
        pdf.ln(1)


def markdown_to_pdf(content: str) -> bytes:
    from fpdf import FPDF

    pdf = FPDF(orientation="P", unit="mm", format="A4")
    pdf.add_page()
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.set_margins(15, 15, 15)

    lines = content.split("\n")
    i = 0
    while i < len(lines):
        s = lines[i].strip()
        if s.startswith("|"):
            table_lines, i = _collect_table_block(lines, i)
            _render_pdf_table(pdf, table_lines)
            continue
        _render_pdf_line(pdf, s)
        i += 1

    return bytes(pdf.output())
